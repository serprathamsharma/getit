"""Resume text extraction and cleaning service supporting PDF, DOCX, and TXT files."""

import io
import logging
import re
from typing import Tuple
from pypdf import PdfReader
import docx

logger = logging.getLogger(__name__)


class ResumeParserService:
    """Handles text extraction and normalization across document formats."""

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted resume text."""
        if not text:
            return ""
        # Remove null characters
        text = text.replace("\x00", "")
        # Replace multiple spaces/newlines with single equivalents
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n+", "\n\n", text)
        return text.strip()

    @classmethod
    def extract_from_pdf(cls, file_bytes: bytes) -> str:
        """Extract text content from a PDF file."""
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            extracted_pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    extracted_pages.append(page_text)
            full_text = "\n\n".join(extracted_pages)
            return cls.clean_text(full_text)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    @classmethod
    def extract_from_docx(cls, file_bytes: bytes) -> str:
        """Extract text content from a DOCX file."""
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text inside tables if present
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        table_texts.append(" | ".join(row_data))
            
            full_text = "\n".join(paragraphs + table_texts)
            return cls.clean_text(full_text)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    @classmethod
    def extract_from_txt(cls, file_bytes: bytes) -> str:
        """Extract text content from a plain text file."""
        try:
            # Try UTF-8 decoding first, fallback to latin-1
            try:
                raw_text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                raw_text = file_bytes.decode("latin-1")
            return cls.clean_text(raw_text)
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            raise ValueError(f"Failed to parse text document: {str(e)}")

    @classmethod
    def parse_file(cls, filename: str, file_bytes: bytes) -> Tuple[str, str]:
        """
        Parse file by extension.
        Returns tuple of (extracted_text, detected_format).
        """
        lower_name = filename.lower()
        if lower_name.endswith(".pdf"):
            return cls.extract_from_pdf(file_bytes), "pdf"
        elif lower_name.endswith(".docx") or lower_name.endswith(".doc"):
            return cls.extract_from_docx(file_bytes), "docx"
        elif lower_name.endswith(".txt") or lower_name.endswith(".md"):
            return cls.extract_from_txt(file_bytes), "txt"
        else:
            # Attempt plain text decoding as fallback
            try:
                return cls.extract_from_txt(file_bytes), "txt"
            except Exception:
                raise ValueError(f"Unsupported file format for '{filename}'. Supported formats: PDF, DOCX, TXT.")
